import sys
import os
import json
import docx

def replace_paragraph_text(p, new_text):
    if p.runs:
        p.runs[0].text = new_text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.text = new_text

def tailor_docx(input_path, output_path, tailored_data):
    doc = docx.Document(input_path)
    
    # 1. Update summary paragraph (P3)
    if len(doc.paragraphs) > 3:
        replace_paragraph_text(doc.paragraphs[3], tailored_data.get('summary', ''))
        
    # 2. Update bullet points
    # Mapping of experience company index to list of paragraph indices
    bullet_mapping = [
        (0, [6, 7, 8, 9]),     # Foundation
        (1, [11, 12]),         # MC Research
        (2, [14, 15, 16, 17]), # Spenmo
        (3, [19, 20]),         # Empala
        (4, [22, 23]),         # Vincere
        (5, [25, 26]),         # Navigos
        (6, [29, 30, 31, 32]), # Paymentwall
        (7, [34])              # KPMG
    ]
    
    experience_data = tailored_data.get('experience', [])
    
    for comp_idx, p_indices in bullet_mapping:
        if comp_idx < len(experience_data):
            comp_data = experience_data[comp_idx]
            bullets = comp_data.get('bullets', [])
            for i, p_idx in enumerate(p_indices):
                if i < len(bullets) and p_idx < len(doc.paragraphs):
                    replace_paragraph_text(doc.paragraphs[p_idx], bullets[i])
                    
    doc.save(output_path)
    print(f"Successfully tailored docx saved at: {output_path}")

if __name__ == "__main__":
    if len(sys.argv) < 4:
        print("Usage: python3 tailor_docx.py <input_docx> <output_docx> <tailored_json_path>")
        sys.exit(1)
        
    input_docx = sys.argv[1]
    output_docx = sys.argv[2]
    json_path = sys.argv[3]
    
    with open(json_path, 'r', encoding='utf-8') as f:
        tailored_data = json.load(f)
        
    tailor_docx(input_docx, output_docx, tailored_data)
